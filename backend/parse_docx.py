from docx import Document

file_path = "data/template.docx"

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)
